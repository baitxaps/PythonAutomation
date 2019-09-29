#! /usr/bin/env python3
# coding: utf-8
import PyPDF2
import os
import docx


# 从 PDF 提取文本
def readpdfFile():
    pdfFileObj = open('./automate_online-materials/meetingminutes.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
    pages = pdfReader.numPages
    print(pages)
    for i in range(pages - 1):
        pageObj = pdfReader.getPage(i)
        print(pageObj.extractText())


#readpdfFile()

# 解密 PDF
def encryptedpefFile():
    pdfReader = PyPDF2.PdfFileReader(open('./automate_online-materials/encrypted.pdf', 'rb'))
    print(pdfReader.isEncrypted)
    pdfReader.decrypt('rosebud')
    pageObj = pdfReader.getPage(0)
    print(pageObj)


#encryptedpefFile()


# 拷贝页面
def copypdffile():
    pdf1File = open('./automate_online-materials/meetingminutes.pdf', 'rb')
    pdf2File = open('/automate_online-materials/meetingminutes2.pdf', 'rb')
    pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
    pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
    pdfWriter = PyPDF2.PdfFileWriter()
    for pageNum in range(pdf1Reader.numPages):
        pageObj = pdf1Reader.getPage(pageNum)
        pdfWriter.addPage(pageObj)

    for pageNum in range(pdf2Reader.numPages):
        pageObj = pdf2Reader.getPage(pageNum)
        pdfWriter.addPage(pageObj)

    pdfOutputFile = open('combinedminutes.pdf', 'wb')
    pdfWriter.write(pdfOutputFile)
    pdfOutputFile.close()
    pdf1File.close()
    pdf2File.close()


#copypdffile()


def pdfRoate():
    minutesFile = open('./automate_online-materials/meetingminutes.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(minutesFile)
    page = pdfReader.getPage(0)
    page.rotateClockwise(90)
    pdfWriter = PyPDF2.PdfFileWriter()
    pdfWriter.addPage(page)
    resultPdfFile = open('rotatedPage.pdf', 'wb')
    pdfWriter.write(resultPdfFile)
    resultPdfFile.close()
    minutesFile.close()


#pdfRoate()


def pdfmergePage():
    minutesFile = open('./automate_online-materials/meetingminutes.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(minutesFile)
    minutesFirstPage = pdfReader.getPage(0)
    pdfWatermarkReader = PyPDF2.PdfFileReader(open('./automate_online-materials/watermark.pdf', 'rb'))
    minutesFirstPage.mergePage(pdfWatermarkReader.getPage(0))
    pdfWriter = PyPDF2.PdfFileWriter()
    pdfWriter.addPage(minutesFirstPage)

    for pageNum in range(1, pdfReader.numPages):
        pageObj = pdfReader.getPage(pageNum)
        pdfWriter.addPage(pageObj)
    resultPdfFile = open('watermarkedCover.pdf', 'wb')
    pdfWriter.write(resultPdfFile)
    minutesFile.close()
    resultPdfFile.close()


#pdfmergePage()


def pdfEncrypted():
    pdfFile = open('./automate_online-materials/meetingminutes.pdf', 'rb')
    pdfReader = PyPDF2.PdfFileReader(pdfFile)
    pdfWriter = PyPDF2.PdfFileWriter()
    for pageNum in range(pdfReader.numPages):
        pdfWriter.addPage(pdfReader.getPage(pageNum))

    pdfWriter.encrypt('swordfish')
    resultPdf = open('encryptedminutes.pdf', 'wb')
    pdfWriter.write(resultPdf)
    resultPdf.close()


#pdfEncrypted()  


#从多个 PDF 中合并选择的页面
def muchPdfMerge():
    pdfFiles = []
    for filename in os.listdir('.'):
        if filename.endswith('.pdf'):
            pdfFiles.append(filename)
    #pdfFiles.sort(key/str.lower)
    pdfWriter = PyPDF2.PdfFileWriter()

    # Loop through all the PDF files.
    for filename in pdfFiles:
        pdfFileObj = open(filename, 'rb')
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        for pageNum in range(1, pdfReader.numPages):
            pageObj = pdfReader.getPage(pageNum)
            pdfWriter.addPage(pageObj)

    # Save the resulting PDF to a file.           
    pdfOutput = open('allminutes.pdf', 'wb')
    pdfWriter.write(pdfOutput)
    pdfOutput.close()


#muchPdfMerge()


def wordDocument():
    doc = docx.Document('./automate_online-materials/demo.docx')
    lens1 = len(doc.paragraphs)
    text1 = doc.paragraphs[0].text
    text2 = doc.paragraphs[1].text
    len2 = len(doc.paragraphs[1].runs)
    text3 = doc.paragraphs[1].runs[0].text
    text4 = doc.paragraphs[1].runs[1].text
    text5 = doc.paragraphs[1].runs[2].text
    text6 = doc.paragraphs[1].runs[3].text
    print('len1: ' + str(lens1) + 'len2: ' + str(len2) + text1 + text2 + text3 + text4 + text5 + text6)


#wordDocument()   


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    doc.add_paragraph('Hello world!')
    doc.add_paragraph('Hello world!', 'Title')
    doc.add_heading('Header 0', 0)  # 标题层次
    doc.add_heading('Header 1', 1)
    doc.add_heading('Header 2', 2)
    doc.add_heading('Header 3', 3)
    doc.add_heading('Header 4', 4)
    doc.paragraphs[0].runs[0].add_break(docx.text.run.WD_BREAK.PAGE)  # 添加换行符和换页符 docx.enum.text.WD_BREAK.PAGE
    doc.add_picture('./automate_online-materials/zophie.png', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
    doc.save('multipleParagraphs.docx')
   
    return '\n'.join(fullText)


print(getText('./automate_online-materials/demo.docx'))






